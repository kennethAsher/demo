#  需要使用的接口请参照：https://github.com/Chyroc/WechatSogou
#  python-docx用于python操作word  传送门：https://python-docx.readthedocs.io/

import docx
from datetime import *
import wechatsogou
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from docx.enum.dml import MSO_THEME_COLOR_INDEX


#用于在word文档中添加超链接
def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url,docx.opc.contants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'),r_id,)

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append(hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

# 文章爬取
def get_articles(headline=True, original=True, timedel=1, add_account=None):
    with open('gzh.txt', 'r') as f:
        accounts = [account.strip() for account in f.readlines()]
        #add_account 必须是一个list或者None
        if add_account is not None:
            if isinstance(list, add_account):
                accounts.extend(add_account)
                with open('gzh.txt', 'w') as w:
                    for account in accounts:
                        f.write(account)
        else:
            print('add_account should be a list')
    ws_api = wechatsogou.WechatSogouAPI(captcha_break_time=3)
    articles = []
    for account in accounts:
        # articles.extend(reformat(ws_api.get_gzh_article_by_history(account)))
        articles.extend(ws_api.get_gzh_article_by_history(account))

    #时间过滤，只选取规定天数以内的
    timestamp = int((datetime.now() - timedelta(days = timedel)).timestamp())
    articles = [article for article in articles if article['datetime']>timestamp]

    # 头条文章过滤，是否选取头条文章，默认是
    if headline:
        articles = [article for article in articles if article['main'] == 1]

    # 原创文章过滤，是否选取原创文章，默认是
    articles = [article for article in articles if article['copyright_stat'] == 100]

    return articles

# 文章整合为文本
def to_mxdocx(data):
    document = Document()
    header = f"公众号新文章（{datetime.now().strftime('%a, %b %d %H:%M')} "
    document.add_heading(header, 0)
    for article in data:
        document.add_paragraph(article['title'], style='ListNumber')
        document.add_paragraph('摘要：'+article['abstract'])
        p = document.add_paragraph('链接：')
        add_hyperlink(p, '小鼠标点击这里哦', article['content_url'])
        document.add_paragraph('来自：'+article['wechat_name'] + '\n')
    p = document.add_paragraph('今天也要元气满满哦~')
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # document.save('/Users/apple/Desktop/{}.docx'.format(header))
    document.save('1.docx')

if __name__ == "__main__":
    articles = get_articles(timedel=1)
    to_msdocx(articles)